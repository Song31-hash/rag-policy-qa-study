"""
샘플 정책 문서 생성. 실험 전 한 번 실행하여 data/policy.docx 를 만듭니다.
"""
from pathlib import Path

def main():
    try:
        from docx import Document
    except ImportError:
        print("pip install python-docx 필요")
        return
    data_dir = Path(__file__).resolve().parent.parent / "data" / "raw"
    data_dir.mkdir(parents=True, exist_ok=True)
    path = data_dir / "policy.docx"
    doc = Document()
    doc.add_paragraph("인사 정책 요약")
    doc.add_paragraph("")
    doc.add_paragraph(
        "휴가 신청은 사용일 2주 전까지 제출해야 합니다. "
        "급한 휴가의 경우 부서장 승인 후 인사팀에 제출할 수 있습니다."
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "원격 근무는 주당 최대 2일까지 가능하며, 사전에 팀장 승인을 받아야 합니다. "
        "재택 근무일에는 업무 연락이 가능한 상태를 유지해야 합니다."
    )
    doc.save(path)
    print(f"Created {path}")

if __name__ == "__main__":
    main()
